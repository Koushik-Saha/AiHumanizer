# humanizer/export_utils.py

import io
from docx import Document
from reportlab.pdfgen import canvas

def generate_docx(submission):
    """
    Generate a DOCX file for a submission.
    Returns an in-memory bytes buffer.
    """
    document = Document()
    document.add_heading(f"Submission #{submission.id}", level=1)
    document.add_paragraph("Original content:")
    document.add_paragraph(submission.original_content)
    document.add_paragraph("Humanized content:")
    document.add_paragraph(submission.humanized_content)
    buf = io.BytesIO()
    document.save(buf)
    buf.seek(0)
    return buf

def generate_pdf(submission):
    """
    Generate a PDF file for a submission.
    Returns an in-memory bytes buffer.
    """
    buf = io.BytesIO()
    p = canvas.Canvas(buf)
    p.setFont("Helvetica", 12)
    p.drawString(50, 800, f"Submission #{submission.id}")
    text = p.beginText(50, 780)
    text.textLines([
        "Original content:",
        submission.original_content,
        "",
        "Humanized content:",
        submission.humanized_content,
    ])
    p.drawText(text)
    p.showPage()
    p.save()
    buf.seek(0)
    return buf